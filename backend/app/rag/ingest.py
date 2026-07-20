

import csv
import io
import logging

from pypdf import PdfReader
import docx
import openpyxl

from app.rag.chunking import chunk_text
from app.rag.embeddings import embed_texts
from app.rag.vectorstore import upsert_chunks, ensure_collection

logger = logging.getLogger("rag.ingest")


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text]
    # Tables aren't walked by `.paragraphs` - pull their cell text too, since
    # a lot of real-world DOCX content (specs, comparisons) lives in tables.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_csv(file_bytes: bytes) -> str:
    # v1: join rows into plain text lines and feed the existing chunker,
    # same as any other document. Row-aware/structured retrieval (treating
    # each row as its own citable unit) is a reasonable future refinement,
    # not required to make CSVs searchable today.
    text = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    lines = [", ".join(row) for row in reader if any(cell.strip() for cell in row)]
    return "\n".join(lines)


def _extract_xlsx(file_bytes: bytes) -> str:
    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(", ".join(cells))
    return "\n".join(parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)
    if ext == "csv":
        return _extract_csv(file_bytes)
    if ext == "xlsx":
        return _extract_xlsx(file_bytes)
    # txt/md fallback
    return file_bytes.decode("utf-8", errors="ignore")


def ingest_document_sync(file_bytes: bytes, filename: str, owner_id: int, document_id: str) -> int:
    """
    Core pipeline: parse -> chunk -> embed -> store in Qdrant.
    Returns number_of_chunks. Raises on failure (caller decides how to record it).
    """
    ensure_collection()

    text = extract_text(file_bytes, filename)
    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("No extractable text found in document")

    vectors = embed_texts(chunks)
    upsert_chunks(document_id, owner_id, filename, chunks, vectors)

    return len(chunks)


def process_document_background(file_bytes: bytes, filename: str, owner_id: int, document_id: str):
    """
    Runs in a FastAPI BackgroundTask, after the HTTP response has already
    been sent - so uploads feel instant to the user instead of blocking on
    embedding. Uses its own DB session since the request-scoped one is gone
    by the time this runs.
    """
    from app.db import SessionLocal, Document  # local import to avoid circular import

    db = SessionLocal()
    try:
        num_chunks = ingest_document_sync(file_bytes, filename, owner_id, document_id)
        doc = db.query(Document).filter(Document.id == document_id).first()
        if doc:
            doc.chunks = num_chunks
            doc.status = "ready"
            db.commit()
    except Exception as e:
        logger.error(f"Document processing failed for {document_id}: {e}")
        doc = db.query(Document).filter(Document.id == document_id).first()
        if doc:
            doc.status = "failed"
            doc.error_message = str(e)[:500]
            db.commit()
    finally:
        db.close()
